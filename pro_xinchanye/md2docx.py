#!/usr/bin/env python3
"""Convert markdown to docx with tables, code blocks, lists, and formatting."""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT as WD_TABLE_ALIGN
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_formatted_paragraph(doc, text, style=None, bold=False, font_size=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
    return p


def parse_inline(text, paragraph):
    """Parse inline formatting (bold, inline code) and add to paragraph."""
    # Handle bold: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Adjust heading styles
    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            h_style.font.size = Pt(18)
        elif i == 2:
            h_style.font.size = Pt(15)
        elif i == 3:
            h_style.font.size = Pt(13)
        elif i == 4:
            h_style.font.size = Pt(11.5)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = open(md_path, 'r', encoding='utf-8').readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    in_list = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                pf = p.paragraph_format
                pf.left_indent = Cm(1)
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                # Add shading
                pPr = p._element.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F5F5F5',
                    qn('w:val'): 'clear',
                })
                pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            if in_list:
                in_list = False
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Headings
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            # Handle nested bold in blockquote
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            i += 1
            continue

        # Table detection: look for | at start and separator row
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Check if this is a table — collect all consecutive table lines
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                table_lines.append(lines[j].strip())
                j += 1

            # Filter out separator rows (|----|----|)
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                # Skip separator rows
                if all(re.match(r'^[-:]+$', c.replace(' ', '')) for c in cells):
                    continue
                data_rows.append(cells)

            if len(data_rows) >= 2:
                # Has header + data — create a real table
                num_cols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=num_cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGN.CENTER

                for ri, row_data in enumerate(data_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = table.cell(ri, ci)
                            # Clean formatting markers from cell
                            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)
                            clean_text = re.sub(r'<br>', '\n', clean_text)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(clean_text)
                            if ri == 0:
                                run.bold = True
                            run.font.size = Pt(8.5)
                            run.font.name = '微软雅黑'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                i = j
                continue
            else:
                # Single row table or just pipe text — treat as normal text
                pass

        # Bullet list
        if re.match(r'^(\s*[-*]\s|\s*\d+\.\s)', line):
            text = re.sub(r'^(\s*[-*]\s|\s*\d+\.\s)', '', line).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        # Remove bold markers for simple paragraphs (inline parsing is complex)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'<br>', '\n', text)
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python3 md2docx.py <markdown_file> [markdown_file...]')
        sys.exit(1)

    for md_file in sys.argv[1:]:
        if md_file.endswith('.md'):
            docx_file = md_file.replace('.md', '.docx')
            print(f'Converting: {md_file} -> {docx_file}')
            convert_md_to_docx(md_file, docx_file)
        else:
            print(f'Skipping non-markdown file: {md_file}')
